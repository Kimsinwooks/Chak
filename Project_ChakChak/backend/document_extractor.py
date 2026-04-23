import os
import json
import zipfile
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from typing import Optional

# ==========================================
# 1. 데이터 구조 및 입력 추상화
# ==========================================

@dataclass
class NormalizedDocument:
    """
    모든 입력(일반 텍스트, 다양한 파일)이 최종적으로 변환되는 표준 텍스트 데이터 구조입니다.
    이후 도메인 추정기(DomainPredictor)는 파일 형식을 모르고 이 객체의 text 속성만 보고 판단합니다.
    """
    text: str
    source_type: str  # 'text', '.pdf', '.docx' 등
    file_path: Optional[str] = None

class InputSource:
    """
    입력이 일반 텍스트인지, 파일인지 판별하고 추상화하는 클래스입니다.
    """
    def __init__(self, data: str, is_file: bool = False):
        self.data = data
        self.is_file = is_file
        
        if self.is_file and not os.path.exists(self.data):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {self.data}")

# ==========================================
# 2. 포맷별 문서 추출기 (Extractors)
# ==========================================

class BaseExtractor:
    def extract(self, file_path: str) -> str:
        raise NotImplementedError

class TxtExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()

class JsonExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return json.dumps(data, ensure_ascii=False)

class PdfExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        import pypdf 
        text = ""
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text.strip()

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        import docx
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

class DocExtractor(BaseExtractor):
    def extract(self, file_path: str) -> str:
        """LibreOffice(soffice)를 이용해 .doc를 임시 폴더에서 .docx로 변환 후 추출"""
        with tempfile.TemporaryDirectory() as temp_dir:
            try:
                subprocess.run([
                    'soffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, file_path
                ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
                base_name = os.path.basename(file_path)
                docx_filename = os.path.splitext(base_name)[0] + ".docx"
                converted_path = os.path.join(temp_dir, docx_filename)
                
                if os.path.exists(converted_path):
                    return DocxExtractor().extract(converted_path)
                else:
                    return ""
            except Exception as e:
                return f"[Error] .doc 변환 실패 (LibreOffice 확인 필요): {str(e)}"

class HwpxExtractor(BaseExtractor):
    @staticmethod
    def check_hwp_com_registration() -> bool:
        import winreg
        try:
            key = winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, r"HWPFrame.HwpObject\CLSID")
            winreg.CloseKey(key)
            print("✅ 한글 COM이 레지스트리에 등록되어 있습니다.")
            return True
        except FileNotFoundError:
            print("📎 HWP COM 등록이 안 되어 있어 등록 시도 중...")
            hwp_exe_path = r"C:\Program Files (x86)\Hnc\Office 2024\HOffice130\Bin\Hwp.exe"
            if os.path.exists(hwp_exe_path):
                subprocess.run([hwp_exe_path, "/regserver"], check=True)
                print("✅ 등록 완료!")
                return True
            else:
                print("❌ Hwp.exe 경로를 찾을 수 없습니다.")
                return False

    @staticmethod
    def write_code_to_manifest(hwpx_path: str):
        import win32com.client
        import hashlib
        
        hwp = None
        try:
            hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
            hwp.Open(os.path.abspath(hwpx_path))
            
            text = hwp.GetTextFile("TEXT", "")
            if text:
                text_bytes = text.encode('utf-8')
                hash_result = hashlib.sha256(text_bytes).hexdigest()
                
                out_path = os.path.join(os.getcwd(), "hashing.txt")
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(hash_result)
        except Exception as e:
            print(f"❌ 해시 추출 중 오류 발생: {e}")
        finally:
            if hwp:
                hwp.Quit()

    def extract(self, file_path: str) -> str:
        """.hwpx 내부 XML을 파싱하여 텍스트 추출 (문단/띄어쓰기 보존)"""
        paragraphs = []
        with zipfile.ZipFile(file_path, 'r') as zf:
            for filename in zf.namelist():
                if filename.startswith("Contents/section") and filename.endswith(".xml"):
                    with zf.open(filename) as f:
                        tree = ET.parse(f)
                        root = tree.getroot()
                        
                        for p_elem in root.iter():
                            if p_elem.tag.endswith('}p'):
                                para_text = []
                                for t_elem in p_elem.iter():
                                    if t_elem.tag.endswith('}t') and t_elem.text:
                                        para_text.append(t_elem.text)
                                
                                if para_text:
                                    paragraphs.append("".join(para_text))
        return "\n".join(paragraphs)

class HwpExtractor(BaseExtractor):
    def __init__(self):
        self.hwpx_extractor = HwpxExtractor()

    def convert_to_hwpx(self, hwp_path: str) -> str:
        """.hwp 파일을 .hwpx 파일로 변환"""
        import win32com.client
        
        hwpx_path = hwp_path + "x"
        if os.path.exists(hwpx_path):
            return hwpx_path

        hwp = None
        try:
            hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
            hwp.Open(os.path.abspath(hwp_path))
            hwp.SaveAs(os.path.abspath(hwpx_path), "HWPX")
        except Exception as e:
            print(f"❌ HWPX 변환 중 오류 발생: {e}")
            raise
        finally:
            if hwp:
                hwp.Quit()
                
        return hwpx_path

    def extract(self, file_path: str) -> str:
        """HWP를 HWPX로 변환한 뒤 추출"""
        hwpx_path = self.convert_to_hwpx(file_path)
        return self.hwpx_extractor.extract(hwpx_path)

# ==========================================
# 3. 문서 추출 계층 (Universal Extractor)
# ==========================================

class UniversalDocumentExtractor:
    """
    InputSource를 받아 파일 확장자에 맞는 Extractor를 매핑하여
    최종적으로 NormalizedDocument 객체로 반환하는 클래스입니다.
    """
    def __init__(self):
        self.extractors = {
            '.txt': TxtExtractor(),
            '.json': JsonExtractor(),
            '.pdf': PdfExtractor(),
            '.docx': DocxExtractor(),
            '.doc': DocExtractor(),
            '.hwpx': HwpxExtractor(),
            '.hwp': HwpExtractor(),  # 주석 해제 및 연결 완료
        }

    def process(self, source: InputSource) -> NormalizedDocument:
        # 1. 파일이 아닌 직접 입력된 텍스트인 경우
        if not source.is_file:
            return NormalizedDocument(text=source.data, source_type="text")
        
        # 2. 파일인 경우 포맷별 추출
        _, ext = os.path.splitext(source.data)
        ext = ext.lower()
        
        if ext in self.extractors:
            extracted_text = self.extractors[ext].extract(source.data)
            return NormalizedDocument(text=extracted_text, source_type=ext, file_path=source.data)
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

# ==========================================
# 실행 테스트 코드
# ==========================================
if __name__ == "__main__":
    extractor = UniversalDocumentExtractor()
    
    # 텍스트 입력 테스트
    # text_input = InputSource("A기관이 10억을 투자하기로 했다.", is_file=False)
    # doc_text = extractor.process(text_input)
    # print("텍스트 추출 결과:", doc_text.text)
    
    # 파일 입력 테스트 로직 (실제 파일이 없으므로 주석 처리)
    # file_input = InputSource("sample_meeting.hwp", is_file=True)
    # doc_file = extractor.process(file_input)
    # print("파일 추출 결과 길이:", doc_file.text)